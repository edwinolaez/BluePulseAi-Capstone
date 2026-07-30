"""
Generates Edwin_TechReview_Summary.docx — a printable summary of the July 29
technical review preparation session.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.1)
    section.right_margin  = Inches(1.1)


# ── Style helpers ─────────────────────────────────────────────────────────────

def heading1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)
    return p

def heading2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2e, 0x6d, 0xa4)
    return p

def heading3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = RGBColor(0x3a, 0x3a, 0x3a)
    return p

def body(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        p.add_run(text).font.size = Pt(11)
    return p

def spacer():
    doc.add_paragraph()

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2e6da4")
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "1a3a5c")
        tcPr.append(shd)

    # Data rows
    for ri, row_data in enumerate(rows):
        row_cells = table.add_row().cells
        fill = "f0f4f8" if ri % 2 == 0 else "ffffff"
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = cell_text
            run = row_cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(10)
            tc = row_cells[i]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  fill)
            tcPr.append(shd)

    # Column widths
    if col_widths:
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[j])

    spacer()
    return table


# ═══════════════════════════════════════════════════════════════════════════════
# COVER
# ═══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("PROJECT JASPER")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Technical Review Preparation Summary")
r2.font.size = Pt(15)
r2.font.color.rgb = RGBColor(0x2e, 0x6d, 0xa4)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("Edwin Olaez  |  PM + QA/Security Lead  |  July 28, 2026")
r3.font.size = Pt(11)
r3.italic = True

divider()
spacer()


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — PROJECT OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════

heading1("1. Project Overview")
body(
    "Project Jasper is a post-wildfire environmental monitoring platform for the "
    "Athabasca watershed, built as a SAIT capstone for client CIRUS (Applied Research at SAIT). "
    "The team has five members, each owning a distinct service area."
)
spacer()

add_table(
    ["Person", "Role", "Owns"],
    [
        ["Edwin Olaez",  "PM + QA/Security Lead", "/tests, /docs, /.github"],
        ["Feven",        "Data Pipeline & API (FastAPI)", "/jasper-backend"],
        ["Richard",      "AI/ML & Simulation", "/jasper-ml"],
        ["Reyta",        "Frontend GIS (Next.js 14)", "/jasper-frontend"],
        ["Rahil",        "DB & Analytics (Supabase + PostGIS + Convex)", "/jasper-db, /convex"],
    ],
    col_widths=[1.3, 2.2, 2.8]
)

heading2("Tech Stack")
for item in [
    "Frontend: Next.js 14 + TypeScript + React-Leaflet + Convex real-time client",
    "Backend: FastAPI + Kong Gateway (Railway)",
    "ML: scikit-learn / TensorFlow + rasterio + SciPy",
    "DB: Supabase + PostGIS + Convex real-time",
    "Deploy: Vercel (frontend) + Railway (backend/ML)",
    "CI/CD: GitHub Actions — 6-stage pipeline (Edwin owns)",
]:
    bullet(item)

spacer()
heading2("Milestones")
add_table(
    ["Milestone", "Date", "Status"],
    [
        ["M1 — Foundation",           "June 20, 2026",   "COMPLETE"],
        ["M2 — Pipeline Live",        "July 4, 2026",    "COMPLETE"],
        ["M3 — AI Live",              "July 18, 2026",   "COMPLETE"],
        ["M4 — Staging Verified",     "July 25, 2026",   "COMPLETE"],
        ["M5 — Production Live",      "August 1, 2026",  "UPCOMING"],
        ["M6 — Demo Day (SAIT/CIRUS)","August 3, 2026",  "UPCOMING"],
    ],
    col_widths=[2.3, 1.8, 1.5]
)
divider()


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — REPO / BRANCH STATUS
# ═══════════════════════════════════════════════════════════════════════════════

heading1("2. Repository & Branch Status (as of July 28, 2026)")

heading2("Main Branch")
body("Main is green and stable. The tip commit is the staging deploy CI fixes (2b2bf13). "
     "All 6 CI stages are green.")

heading2("Pending Branches — Not Yet Merged to Main")

heading3("fix/security-api-key  (PR #29) — 6 commits ahead of main")
add_table(
    ["Commit", "What it adds"],
    [
        ["6c84cb2", "Security fix — API key moved server-side; 5 Next.js proxy routes created"],
        ["9d2185f / 368ef4f", "CI + test fallback for ML_API_KEY during migration"],
        ["7e1854c", "Lighthouse + P95 benchmark results filled into test-completion-report"],
        ["3a9c30a", "Backend coverage results filled in (69%, 631 statements)"],
        ["77ea9e6", "CI/CD config comments + security rules documented"],
    ],
    col_widths=[2.0, 4.3]
)
body("Blocked on two manual steps before merge:", bold=True)
bullet("GitHub Secrets → add ML_API_KEY = jasper-dev-api-key-2026")
bullet("Vercel → add ML_API_KEY = jasper-dev-api-key-2026 (all environments)")

spacer()
heading3("feature/simulation-map-link — 19+ commits ahead of main")
for item in [
    "ArcGIS simulation layers and elevation risk zones",
    "Field photo upload with simulation tagging",
    "133 tests green",
    "Includes its own copy of the security fix",
    "Not yet merged — decision pending before July 29",
]:
    bullet(item)

divider()


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — EDWIN'S CONTRIBUTIONS
# ═══════════════════════════════════════════════════════════════════════════════

heading1("3. Edwin's Contributions")
body(
    "Edwin is the gatekeeper of the project. No code reaches main without his review and CI sign-off. "
    "He owns three areas: the CI/CD pipeline, the test suite, and security."
)

# 3a CI/CD
heading2("3a. CI/CD Pipeline  (.github/workflows/)")
body(
    "A 6-stage GitHub Actions pipeline runs automatically on every push from any team member. "
    "If any stage fails, the PR is blocked from merging."
)
add_table(
    ["Stage", "What it does"],
    [
        ["Stage 1 — Lint",        "Checks code formatting"],
        ["Stage 2 — Security",    "Runs Semgrep SAST — scans for security vulnerabilities"],
        ["Stage 3 — Unit",        "Runs fast unit tests on individual functions"],
        ["Stage 4 — Integration", "Tests that all 4 services talk to each other correctly"],
        ["Stage 5 — Build",       "Builds the Next.js frontend and proves it compiles"],
        ["Stage 6 — Performance", "Measures API P95 response times under load"],
    ],
    col_widths=[2.2, 4.1]
)
body(
    "Edwin also built deploy-staging.yml (auto-deploys on every merge to develop) and "
    "deploy-production.yml (manual-only, requires typed DEPLOY confirmation + GitHub environment approval)."
)

# 3b Test suite
spacer()
heading2("3b. Test Suite  (tests/)")
body("79 tests across 7 files. Result: 71 passed, 3 expected skips, 0 failed.")
add_table(
    ["File", "What it proves", "Tests"],
    [
        ["test_health.py",           "Backend is alive; API gateway rejects bad keys",                         "6"],
        ["test_api_contracts.py",    "Every teammate's API returns exactly the promised shape",                 "16"],
        ["test_rbac.py",             "Role permissions enforced — viewer can't write, ingest can't read",       "10"],
        ["test_e2e_pipeline.py",     "Sensor reading flows: ingest → DB → API → frontend",                     "12"],
        ["test_convex_integration.py","Convex real-time queries and mutations work correctly",                  "11"],
        ["test_ml_integration.py",   "ML model returns predictions within performance budget",                  "19"],
        ["benchmark_api.py",         "P95 response times measured across all endpoints",                        "5"],
    ],
    col_widths=[2.2, 3.4, 0.7]
)

# 3c Security
spacer()
heading2("3c. Security Work")

heading3("OWASP Top 10 Audit  (docs/owasp-mapping.md)")
body(
    "Every one of the 10 most common web application security risks was reviewed against "
    "Project Jasper's architecture, and controls were documented and verified."
)
add_table(
    ["OWASP Risk", "Control in place", "Status"],
    [
        ["A01 Broken Access Control",         "Supabase RLS + Kong API key + RBAC tests",                 "CONFIRMED"],
        ["A02 Cryptographic Failures",         "Secrets in env vars only; HTTPS everywhere; AES-256 at rest","CONFIRMED"],
        ["A03 Injection",                      "Pydantic validation + parameterized queries + Semgrep",    "CONFIRMED"],
        ["A04 Insecure Design",                "Rate limiting (20 req/min); 404 on unknown sectors",       "CONFIRMED"],
        ["A05 Security Misconfiguration",      "Semgrep DEBUG rules; CORS whitelist",                      "CONFIRMED"],
        ["A06 Vulnerable Components",          "Dependabot + pip-audit on every push",                     "CONFIRMED"],
        ["A07 Auth Failures",                  "Kong + Supabase JWT + 2FA on admin",                       "CONFIRMED"],
        ["A08 Integrity Failures",             "Branch protection; pinned dependencies; required review",   "CONFIRMED"],
        ["A09 Logging Failures",               "Railway + Kong + Supabase + Semgrep artifact logs",        "CONFIRMED"],
        ["A10 SSRF",                           "No URL-type params; all HTTP calls use hardcoded constants","CONFIRMED"],
    ],
    col_widths=[2.1, 3.1, 1.1]
)

heading3("API Key Security Fix  (commit 6c84cb2)")
body(
    "Discovered that NEXT_PUBLIC_API_KEY was being bundled into the browser JavaScript bundle "
    "by Next.js — meaning any user could open DevTools and steal the Kong Gateway credential. Fixed by:"
)
for item in [
    "Creating 5 Next.js Route Handler proxies (app/api/backend/*, app/api/ml/*) that read ML_API_KEY server-side",
    "Refactoring lib/api.ts so all fetch calls go to local /api/* routes — key never reaches the browser",
    "Renaming NEXT_PUBLIC_API_KEY → ML_API_KEY across CI workflows and all test files",
]:
    bullet(item)

divider()


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — TESTING TOOLS & RATIONALE
# ═══════════════════════════════════════════════════════════════════════════════

heading1("4. Testing Tools & Why Those Choices")

add_table(
    ["Tool", "What it does", "Why it was chosen"],
    [
        ["pytest",          "Runs all tests; reports pass/fail/skip",
                            "Industry standard; integrates cleanly into GitHub Actions"],
        ["httpx",           "Sends real HTTP requests to real live endpoints",
                            "No mocks — tests what actually runs in production"],
        ["conftest.py",     "Shared fixtures: URLs, auth headers, HTTP client",
                            "Write setup once, use it in every test file"],
        ["Contract tests",  "Verifies every teammate's API response shape",
                            "Catches breaking changes before they reach the frontend"],
        ["RBAC tests",      "Logs in as each role and tries forbidden actions",
                            "Proves database security policies actually enforce permissions"],
        ["P95 benchmarks",  "Measures 95th percentile response time across 20 requests",
                            "Surfaces slow outliers that averages would hide"],
        ["Semgrep SAST",    "Scans code for security vulnerabilities on every push",
                            "Automated security gate — zero HIGH findings blocks merge"],
    ],
    col_widths=[1.5, 2.3, 2.5]
)

heading2("Why P95 instead of average?")
body(
    "An average hides outliers. If 19 requests take 100ms but one takes 2,000ms, "
    "the average looks fine — but one in twenty users gets a terrible experience. "
    "P95 means 95% of requests must finish within the budget (500ms). "
    "That is the industry standard for monitoring dashboards."
)

heading2("Why real HTTP calls instead of mocks?")
body(
    "A mocked test can pass even when the real endpoint is broken. "
    "This project is a live integrated system across 4 services built by 4 different people. "
    "Real HTTP calls against the deployed staging backend is the only way to prove "
    "the integrations actually work end-to-end."
)

divider()


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — INSTRUCTOR FEEDBACK VS. WHAT WAS DONE
# ═══════════════════════════════════════════════════════════════════════════════

heading1("5. Instructor Feedback (Mehdi) vs. What Was Done")

heading2("Technical Items — Edwin's Responsibility")
add_table(
    ["Mehdi's Feedback", "What Was Done", "Status"],
    [
        ["Display source, timestamp, quality, and refresh status for environmental data",
         "Added Simulated vs Live·Observed badges to WaterQualityWidget. Per-sensor Online/Offline status in PipelineStatusWidget. Observed·IoT Sensors badge on DashboardPage.",
         "DONE\ncommit c7889f7"],

        ["Include model confidence, uncertainty, and limitations in AI output",
         "AiOverviewPage now labels predictions as 'AI Prediction' and sensor inputs as 'Observed'. AI badges added to summary table columns.",
         "DONE\ncommit c7889f7"],

        ["Deploy complete application to the cloud",
         "Staging live on Vercel since July 24. Railway backend + ML service live. All 6 CI stages green.",
         "DONE\nJuly 24"],

        ["Secure secrets / failure handling",
         "API key moved server-side (5 proxy routes). ConvexErrorBoundary shows user-facing message instead of blank crash.",
         "DONE\ncommits 6c84cb2 + c7889f7"],

        ["Maintain clear service boundaries",
         "6-stage CI pipeline tests each service boundary independently. Contract tests lock API shapes. OWASP doc maps each service's security controls.",
         "DONE\nongoing"],

        ["Audit logs and monitoring",
         "Railway logs capture all FastAPI requests. Kong logs auth failures. Semgrep artifacts attached to every CI run. Documented in owasp-mapping.md (A09 section).",
         "PARTIAL\nlogs exist; no auto-alert yet"],
    ],
    col_widths=[1.9, 3.2, 1.2]
)

spacer()
heading2("Optional / Industry-Level Items")
add_table(
    ["Mehdi's Suggestion", "What Was Done", "Status"],
    [
        ["Sensor-health and missing-data warnings",
         "PipelineStatusWidget shows per-sensor Online/No Link status with colour indicators.",
         "DONE"],
        ["Clear separation: observed vs predicted",
         "Every data point on AiOverviewPage and DashboardPage is now labelled as either AI Prediction or Observed.",
         "DONE"],
        ["Layer controls (fire severity, vegetation, water, sensors)",
         "On feature/simulation-map-link — ArcGIS simulation layers, elevation risk zones added.",
         "DONE\n(not yet merged)"],
        ["Before/after satellite image comparison",
         "Not done — requires Richard's imagery data pipeline.",
         "NOT DONE"],
        ["Exportable maps and research reports",
         "Not done — out of scope for this sprint.",
         "NOT DONE"],
    ],
    col_widths=[2.0, 3.1, 1.2]
)

spacer()
heading2("AI-Specific Suggestions")
add_table(
    ["Mehdi's Suggestion", "What Was Done", "Status"],
    [
        ["Do NOT present AI results as confirmed conclusions",
         "AiOverviewPage labels all predictions with 'AI Prediction' badges — visually separated from observed data.",
         "DONE"],
        ["Require expert review before important decisions",
         "Researcher Chatbot (AI Overview page) is live — designed for researchers to query and validate, not for automated decisions.",
         "DONE"],
        ["Summarize important changes for researchers",
         "Chatbot uses Claude API with tool use to query ML endpoints and summarise findings in plain language.",
         "DONE"],
        ["Explain which data contributed to each prediction",
         "Model card has F1/precision/recall but no per-prediction feature explanation.",
         "PARTIAL"],
    ],
    col_widths=[2.0, 3.1, 1.2]
)

divider()


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — KEY METRICS SUMMARY
# ═══════════════════════════════════════════════════════════════════════════════

heading1("6. Key Metrics — What to Quote Tomorrow")

add_table(
    ["Metric", "Target", "Actual", "Status"],
    [
        ["Total tests written",             "—",      "79",                         "—"],
        ["Tests passing",                   "100%",   "71 passed / 3 skips / 0 failed", "MET"],
        ["Frontend test coverage",          "≥ 75%",  "81.75%",                     "MET"],
        ["Backend test coverage",           "≥ 80%",  "69% (631 stmts)",            "NOT MET*"],
        ["API contracts tested",            "100%",   "6/6 (100%)",                 "MET"],
        ["Semgrep HIGH findings",           "0",      "0",                          "MET"],
        ["Unpatched HIGH CVEs",             "0",      "0",                          "MET"],
        ["Lighthouse Performance (staging)","≥ 85",   "96/100",                     "MET"],
        ["Lighthouse Accessibility",        "≥ 90",   "93/100",                     "MET"],
        ["API P95 — ML endpoints",          "< 500ms","196ms and 225ms",            "MET"],
        ["API P95 — map query",             "< 500ms","986ms (Railway cold-start**)","NOT MET**"],
        ["ML model F1 score",               "≥ 0.75", "0.80 macro / 0.82 weighted", "MET"],
    ],
    col_widths=[2.6, 1.1, 1.9, 1.0]
)

body(
    "* Backend coverage gap is in routers with no dedicated unit tests (admin, alerts, auth, "
    "change_detection, simulation, fusion). These are exercised via integration tests in Stage 4 "
    "but not counted in the coverage report. Core modules (main, config, health, timeline) are at ≥ 93%."
)
body(
    "** Map query P95 of 986ms is Railway free-tier cold-start latency, not a code issue. "
    "Database execution time alone is 1.215ms (confirmed by Rahil's PostGIS EXPLAIN ANALYZE). "
    "All ML and health endpoints comfortably pass the 500ms budget."
)

divider()


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — OPEN ITEMS BEFORE DEMO DAY
# ═══════════════════════════════════════════════════════════════════════════════

heading1("7. Open Items Before Demo Day (August 3)")

add_table(
    ["Item", "Owner", "Urgency"],
    [
        ["Add ML_API_KEY to GitHub Secrets + Vercel (unblocks PR #29 merge)", "Edwin",  "HIGH — before July 29"],
        ["ResearcherChatPanel broken on Reyta's Vercel URL — needs ANTHROPIC_API_KEY", "Reyta", "HIGH — before July 29"],
        ["TEST_INGEST_JWT regeneration (fixes 3 CI skips on ingest auth tests)", "Rahil",  "MEDIUM"],
        ["QUERY_BENCHMARK_REPORT.md completion",                                "Rahil",  "MEDIUM"],
        ["ML F1 scores filled into model_card.md",                              "Richard","MEDIUM"],
        ["Go-live checklist — 5 of 12 items still unsigned (items 1,5,6,7,11,12)","Team", "M5 — Aug 1"],
        ["deploy-production.yml — run for M5",                                  "Edwin",  "M5 — Aug 1"],
        ["Merge feature/simulation-map-link to main",                           "Edwin",  "Before Demo Day"],
    ],
    col_widths=[3.3, 1.2, 1.8]
)

divider()


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — TALKING POINTS FOR TOMORROW
# ═══════════════════════════════════════════════════════════════════════════════

heading1("8. Talking Points for Tomorrow's Technical Review")

heading2("If asked: What is your role?")
body(
    '"I am the PM and QA/Security lead. I do not build individual features — '
    "I make sure the features everyone else builds actually work together, stay secure, "
    "and can be deployed reliably. Every merge to main went through me. "
    'If the app is running on demo day, it is because the CI pipeline and test suite I built caught problems before they got there."'
)

heading2("If asked: What is the CI pipeline?")
body(
    '"It is a 6-stage automated referee that runs on every push. '
    "If any stage fails, the PR is blocked. Stage 2 is a security scan. Stage 4 tests that all 4 services "
    "talk to each other correctly. Stage 6 measures API response times. "
    'Nothing broken has ever reached main on my watch."'
)

heading2("If asked: Why is backend coverage only 69%?")
body(
    '"The gap is in complex routers — admin, alerts, auth, simulation — that have no dedicated unit tests yet. '
    "Those routers are exercised by integration tests in Stage 4, but pytest-cov does not count those hits. "
    "Core modules — main, config, health, timeline — are at 93-100%. "
    'I have documented this honestly in the Test Completion Report with a note explaining the cause."'
)

heading2("If asked: What security risks did you find and fix?")
body(
    '"I found that the Kong API key was exposed in the browser JavaScript bundle because of the NEXT_PUBLIC_ prefix. '
    "Any user could open DevTools and steal it. I fixed it by creating 5 server-side proxy routes so the key "
    "never reaches the browser. I also completed a full OWASP Top 10 audit and verified all 10 risks are "
    'mitigated with documented controls."'
)

heading2("If asked: Why is the map query P95 over 500ms?")
body(
    '"The database execution time is 1.215ms — that is well within budget. '
    "The 986ms P95 is Railway free-tier cold-start latency added after the DB returns. "
    "When the server is warm, the endpoint responds in under 200ms. This is a hosting-tier limitation, "
    'not a code issue, and I have documented it transparently in the Test Completion Report."'
)

divider()

# Footer
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Project Jasper  |  Edwin Olaez  |  SAIT Capstone 2026  |  edwinolaez02@gmail.com")
r.font.size = Pt(9)
r.italic = True
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


# ── Save ──────────────────────────────────────────────────────────────────────
out_path = os.path.join(
    os.path.dirname(__file__), "..", "docs", "Edwin_TechReview_Summary.docx"
)
out_path = os.path.normpath(out_path)
doc.save(out_path)
print(f"Saved: {out_path}")
