"""Generates Edwin's July 29 Technical Review prep guide as a .docx file."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1.1)
section.right_margin = Inches(1.1)

# ── Styles ────────────────────────────────────────────────────────────────────
BLUE  = RGBColor(0x1E, 0x40, 0xAF)   # deep blue
GRAY  = RGBColor(0x4B, 0x55, 0x63)
BLACK = RGBColor(0x11, 0x18, 0x27)

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = BLUE
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK
    return p

def body(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        r.bold = True
        r.font.size = Pt(10)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = BLACK
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    return p

def quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.35)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(f'"{text}"')
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    return p

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            cells[i].paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE BLOCK
# ══════════════════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("Project Jasper — July 29 Technical Review")
tr.bold = True
tr.font.size = Pt(18)
tr.font.color.rgb = BLUE

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("Presentation Prep Guide  |  Edwin Olaez  |  PM + QA/Security Lead")
sr.font.size = Pt(11)
sr.font.color.rgb = GRAY

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
dr = date_p.add_run("Prepared: July 25, 2026  |  SAIT Capstone — BluePulse AI")
dr.font.size = Pt(10)
dr.font.color.rgb = GRAY

doc.add_paragraph()
divider()
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# PART 1 — SYSTEM OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
heading1("Part 1 — System Overview (Your Opening 60 Seconds)")
body("Memorise and deliver this without reading. It sets the stage for the whole team.")
quote(
    "Project Jasper is a post-wildfire environmental monitoring platform for the Athabasca watershed. "
    "It has five integrated services: Feven's FastAPI backend ingests real IoT and satellite data through "
    "a Kong API gateway deployed on Railway. Rahil's Supabase+PostGIS database stores and queries that data "
    "with spatial indexes. Richard's ML service runs three models — burn scar change detection, erosion "
    "simulation, and contaminant spread — all deployed on Railway. Reyta's Next.js frontend renders live GIS "
    "maps, sensor widgets, and AI outputs. Convex provides real-time data sync across the frontend. "
    "My role was to make sure all of it connects — through a 6-stage CI/CD pipeline, security gates, "
    "and a formal test suite."
)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# PART 2 — YOUR CONTRIBUTION
# ══════════════════════════════════════════════════════════════════════════════
heading1("Part 2 — Your Own Contribution (What YOU Built)")
body("Mehdi will ask you directly. Have a specific answer for each area.")

heading2("CI/CD Pipeline  (.github/workflows/)")
bullet("Built and own a 6-stage pipeline: Lint → Security → Unit → Integration → Build → Performance")
bullet("All 6 stages green on develop (SHA e258628, July 24, 2026)")
bullet("Wired in all teammate secrets: SUPABASE_URL, ML_API_URL, RAILWAY_API_URL, ANTHROPIC_API_KEY")
bullet("Personally unblocked Feven (Railway ingest secrets), Richard (ML proxy paths), Rahil (Convex URL)")

heading2("Test Suite  (tests/)")
bullet("79 tests across 7 files — you wrote all of them")
bullet("71 passed, 3 expected skips (ingest JWT), 0 failures")
bullet("Covers: health checks, API contracts, RBAC, E2E pipeline, Convex integration, ML integration, P95 benchmarks")

heading2("Security")
bullet("0 Semgrep HIGH findings, 0 unpatched HIGH CVEs")
bullet("Moved backend API key from NEXT_PUBLIC_API_KEY (browser-visible) to server-side ML_API_KEY — branch: fix/security-api-key")
bullet("RBAC: all 4 roles (admin, analyst, ingest, viewer) verified through Supabase RLS tests")
bullet("OWASP Top 10 mapped in docs/owasp-mapping.md")

heading2("Documentation")
bullet("docs/test-completion-report.md — formal SAIT faculty deliverable, fully completed")
bullet("docs/api-contracts.md — 6 inter-module contracts, all confirmed")
bullet("CI feedback docs delivered to Richard, Feven, and Rahil when their services broke CI")

divider()

# ══════════════════════════════════════════════════════════════════════════════
# PART 3 — YOUR NUMBERS
# ══════════════════════════════════════════════════════════════════════════════
heading1("Part 3 — Your Numbers (Know These Cold)")

add_table(
    ["Metric", "Target", "Actual", "Status"],
    [
        ("Frontend test coverage",      "≥ 75%",    "81.75%",            "PASS ✓"),
        ("Backend test coverage",       "≥ 80%",    "69%",               "MISS — see explanation below"),
        ("API contracts tested",        "100%",     "100% (6/6)",        "PASS ✓"),
        ("Semgrep HIGH findings",       "0",        "0",                 "PASS ✓"),
        ("Unpatched HIGH CVEs",         "0",        "0",                 "PASS ✓"),
        ("Integration tests passing",  "100%",     "96% (71/74)",       "PASS ✓ — 3 expected skips"),
        ("Lighthouse Performance",      "≥ 85",     "96",                "PASS ✓"),
        ("Lighthouse Accessibility",    "≥ 90",     "93",                "PASS ✓"),
        ("API P95 — /health",           "< 200ms",  "176.6ms",           "PASS ✓"),
        ("API P95 — ML endpoints",      "< 500ms",  "~200ms",            "PASS ✓"),
        ("API P95 — map query",         "< 500ms",  "986.6ms",           "MISS — see explanation below"),
        ("ML F1 score",                 "≥ 0.75",   "0.80 macro",        "PASS ✓"),
        ("DB spatial query",            "< 500ms",  "1.215ms execution", "PASS ✓"),
    ]
)

heading2("Explanation for the Two Misses")
body(
    "Backend 69%:",
    bold_prefix="►"
)
body(
    "The core modules — main, config, health, timeline — are all at 93–100%. "
    "The gap is in routers that are only exercised through Stage 4 integration tests, not isolated unit tests. "
    "The integration pipeline covers them, but they don't count toward the unit coverage metric."
)
doc.add_paragraph()
body(
    "Map query P95 986ms:",
    bold_prefix="►"
)
body(
    "The DB itself executes in 1.215ms — that is Rahil's PostGIS benchmark. "
    "The overage is Railway free-tier cold-start latency, not a code issue. On a paid tier this resolves."
)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# PART 4 — HOW YOUR WORK CONNECTS
# ══════════════════════════════════════════════════════════════════════════════
heading1("Part 4 — How Your Work Connects to the Full App")
body("Mehdi wants to hear integration, not just 'I did tests.' Say it like this:")
quote(
    "Every time a teammate's service changed an API contract, my contract tests caught the break before it "
    "hit staging. When Feven's ingest service was returning 500s, my CI feedback doc identified the missing "
    "Supabase env vars within an hour. When Richard's ML endpoints were returning 502s, I traced it to the "
    "Railway deployment logs and documented exactly what he needed to fix. The CI pipeline is how the whole "
    "team's work stays integrated — without it, five independent services would never connect cleanly."
)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# PART 5 — DEMO PLAN
# ══════════════════════════════════════════════════════════════════════════════
heading1("Part 5 — Demo Plan (2–3 Minutes for Your Contribution)")
bullet("Open GitHub Actions → show the 6-stage pipeline green")
bullet("Open docs/test-completion-report.md → walk through the filled-in metrics table")
bullet("Show one test file (tests/test_rbac.py) → explain what RBAC testing means and why it matters for security")
bullet("Show the fix/security-api-key commit → explain what you moved and why (NEXT_PUBLIC_ prefix exposes secrets to the browser)")
doc.add_paragraph()
body("Backup plan:", bold_prefix="►")
body("Have a screen recording of the pipeline running in case GitHub Actions is slow on the day.")

divider()

# ══════════════════════════════════════════════════════════════════════════════
# PART 6 — SOFT SKILLS
# ══════════════════════════════════════════════════════════════════════════════
heading1("Part 6 — Soft Skills (From Mehdi's Written Feedback)")
bullet("Do not read off notes — glance at them, then look up and talk")
bullet("Slow down when citing numbers — pause after each one")
bullet("Make eye contact with Mehdi, not the screen")
bullet("Stay still and attentive when teammates are speaking — no side-talking")
bullet("Equal speaking time across the whole team")

divider()

# ══════════════════════════════════════════════════════════════════════════════
# PART 7 — THE ONE THING TO MEMORISE
# ══════════════════════════════════════════════════════════════════════════════
heading1("Part 7 — The One Answer to Memorise")
body('If Mehdi asks "what was the hardest part of your role?" — answer this:')
quote(
    "The hardest part was keeping five independently-developed services integrated. "
    "Everyone was building in separate branches, with separate tools and separate deployments. "
    "My job was to make sure that when you put it all together, it actually works — "
    "through contracts, tests, and a CI pipeline that catches breaks before they hit staging."
)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# PART 8 — QUICK REFERENCE CARD
# ══════════════════════════════════════════════════════════════════════════════
heading1("Part 8 — Quick Reference Card (Print This Page)")

heading2("Team Roles — One Sentence Each")
add_table(
    ["Person", "Role", "Tech"],
    [
        ("Edwin",   "PM + QA/Security",          "GitHub Actions, pytest, Semgrep"),
        ("Feven",   "Data Pipeline & API",        "FastAPI, Kong, Railway"),
        ("Richard", "AI/ML & Simulation",         "scikit-learn, rasterio, SciPy"),
        ("Reyta",   "Frontend GIS",               "Next.js 14, React-Leaflet, Convex"),
        ("Rahil",   "DB & Analytics",             "Supabase, PostGIS, Convex"),
    ]
)

heading2("Key URLs")
bullet("Staging frontend: https://bluepulseai-capstone-drtxqator-blue-pulse-ai-capstone.vercel.app")
bullet("Railway backend: https://bluepulseai-capstone-production.up.railway.app/health")
bullet("CI runs: https://github.com/edwinolaez/BluePulseAi-Capstone/actions")
bullet("Convex: https://cautious-dogfish-127.convex.cloud")

heading2("Your Key Files to Show")
bullet("tests/test_rbac.py — RBAC security tests")
bullet("tests/benchmark_api.py — P95 performance benchmarks")
bullet(".github/workflows/ci.yml — 6-stage pipeline")
bullet("docs/test-completion-report.md — formal QA deliverable")
bullet("docs/owasp-mapping.md — OWASP Top 10 mapping")

# ══════════════════════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer_p.add_run("Project Jasper  |  SAIT Capstone 2026  |  Edwin Olaez  |  Prepared July 25, 2026")
fr.font.size = Pt(8)
fr.font.color.rgb = GRAY

out = r"C:\Users\Edwin Olaez\Documents\BluePulseAi-Capstone\docs\Edwin_July29_TechReview_Prep.docx"
doc.save(out)
print(f"Saved: {out}")
