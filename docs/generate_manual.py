"""Generates the Project Jasper User Manual as a DOCX file."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.5)

# ── Colour palette ────────────────────────────────────────────────────────────
TEAL   = RGBColor(0x00, 0x87, 0x9B)
DARK   = RGBColor(0x1A, 0x1A, 0x2E)
GREY   = RGBColor(0x55, 0x55, 0x66)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
AMBER  = RGBColor(0xD9, 0x7B, 0x06)
GREEN  = RGBColor(0x15, 0x80, 0x3D)
RED    = RGBColor(0xB9, 0x1C, 0x1C)

# ── Helper functions ──────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Set background colour of a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(18)
    run.font.color.rgb = TEAL
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(13)
    run.font.color.rgb = DARK
    return p

def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(11)
    run.font.color.rgb = GREY
    return p

def body(text, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size   = Pt(11)
    run.italic      = italic
    run.font.color.rgb = DARK
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        rb = p.add_run(bold_prefix + " ")
        rb.bold = True
        rb.font.size = Pt(11)
        rb.font.color.rgb = DARK
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.color.rgb = DARK
    return p

def tip_box(text):
    """A shaded tip paragraph."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_bg(cell, "E8F5F7")
    cell.width = Inches(5.8)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    rb = p.add_run("💡 Tip:  ")
    rb.bold = True
    rb.font.size = Pt(10.5)
    rb.font.color.rgb = TEAL
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = DARK
    doc.add_paragraph()

def note_box(text):
    """A shaded note paragraph."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_bg(cell, "FFF8E7")
    cell.width = Inches(5.8)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    rb = p.add_run("📌 Note:  ")
    rb.bold = True
    rb.font.size = Pt(10.5)
    rb.font.color.rgb = AMBER
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = DARK
    doc.add_paragraph()

def divider():
    p = doc.add_paragraph("─" * 72)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        run.font.size = Pt(8)

def page_break():
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(60)
r = cover.add_run("PROJECT JASPER")
r.bold = True
r.font.size = Pt(32)
r.font.color.rgb = TEAL

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("Athabasca Watershed Monitoring Platform")
r2.font.size = Pt(16)
r2.font.color.rgb = GREY

doc.add_paragraph()

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = tagline.add_run("User Guide for CERCUTS / ARIS Staff")
r3.bold = True
r3.font.size = Pt(14)
r3.font.color.rgb = DARK

doc.add_paragraph()
doc.add_paragraph()

ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = ver.add_run(f"Version 1.0   ·   {datetime.date.today().strftime('%B %Y')}   ·   Prepared by Team BluePulse AI")
r4.font.size = Pt(10)
r4.font.color.rgb = GREY
r4.italic = True

doc.add_paragraph()
divider()
doc.add_paragraph()

intro_note = doc.add_paragraph()
intro_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
r5 = intro_note.add_run(
    "This guide is written for environmental scientists, field coordinators, and\n"
    "watershed managers. No technical background is required."
)
r5.font.size = Pt(11)
r5.font.color.rgb = GREY
r5.italic = True

page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════════════════════

heading1("Table of Contents")
toc_items = [
    ("1.", "Getting Started"),
    ("2.", "Navigating the Application"),
    ("3.", "Map View — Your Main Workspace"),
    ("4.", "Dashboard — Water Quality Analytics"),
    ("5.", "Reports — Download & Generate Reports"),
    ("6.", "Archives — Historical Survey Data"),
    ("7.", "Alerts & What They Mean"),
    ("8.", "Live System Logs"),
    ("9.", "Glossary of Terms"),
    ("10.", "Frequently Asked Questions"),
    ("11.", "Contact & Support"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    rn = p.add_run(f"  {num}  ")
    rn.bold = True
    rn.font.color.rgb = TEAL
    rn.font.size = Pt(11)
    rt = p.add_run(title)
    rt.font.size = Pt(11)
    rt.font.color.rgb = DARK

page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 1 — GETTING STARTED
# ══════════════════════════════════════════════════════════════════════════════

heading1("1.  Getting Started")

body(
    "Project Jasper is a web-based environmental monitoring platform built to help "
    "CERCUTS and ARIS staff track post-wildfire recovery conditions across the "
    "Athabasca watershed in real time. You can view water quality data, monitor "
    "erosion and contamination risk zones on a map, download reports, and browse "
    "historical survey archives — all from your web browser."
)

heading2("What You Need")
bullet("A modern web browser (Google Chrome, Microsoft Edge, or Firefox recommended).")
bullet("An internet connection.")
bullet("Your login credentials, provided by your system administrator.")

tip_box(
    "The platform works best on a desktop or laptop screen. While it can be viewed "
    "on a tablet, some map features may be easier to use on a larger display."
)

heading2("Opening the Application")
body("Your system administrator will provide you with the web address (URL) for the platform. Once you have it:")
bullet("Step 1 —", "Open your web browser.")
bullet("Step 2 —", "Type or paste the web address into the address bar and press Enter.")
bullet("Step 3 —", "The application will load and you will see the main Map View screen.")

note_box(
    "If the page does not load, check your internet connection or contact your "
    "administrator. The platform may occasionally be updated — if you see an "
    "unexpected screen, try refreshing your browser (press F5 or Ctrl+R)."
)

page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 2 — NAVIGATING THE APPLICATION
# ══════════════════════════════════════════════════════════════════════════════

heading1("2.  Navigating the Application")

body(
    "The application has two navigation areas: a top bar with the main section tabs, "
    "and a left sidebar with additional map and monitoring controls."
)

heading2("Top Navigation Bar")
body("At the very top of the screen you will find four tabs:")

tab_table = doc.add_table(rows=5, cols=2)
tab_table.style = "Table Grid"
headers = ["Tab", "What it does"]
for i, h in enumerate(headers):
    cell = tab_table.cell(0, i)
    set_cell_bg(cell, "005F6E")
    p = cell.paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = WHITE
    r.font.size = Pt(11)

rows_data = [
    ("Map View",   "The main interactive map of the watershed. This is where you will spend most of your time — viewing hazard zones, sector statuses, and live water quality readings."),
    ("Dashboard",  "A summary of water quality metrics over time, with charts and alert summaries for all monitoring stations."),
    ("Reports",    "A library of environmental reports you can download as PDFs, or run new model simulations."),
    ("Archives",   "Historical survey snapshots and satellite pass data from previous monitoring periods."),
]
for i, (tab, desc) in enumerate(rows_data, start=1):
    tab_table.cell(i, 0).text = tab
    tab_table.cell(i, 1).text = desc
    tab_table.cell(i, 0).paragraphs[0].runs[0].bold = True

doc.add_paragraph()

heading2("Left Sidebar")
body("The sidebar on the left side of the screen gives you quick access to:")
bullet("Layers & Map —", "Return to the main map view.")
bullet("Analytics —", "Go to the Dashboard.")
bullet("Telemetry Focus —", "See a list of all monitoring sectors with their current status.")
bullet("Legend Panel —", "View what each colour on the map means.")
bullet("Export GeoJSON —", "Download the current map data in a format compatible with GIS software.")
bullet("Support Request —", "Contact the support team.")
bullet("Diagnostic Logs —", "Open the Live System Logs panel.")

heading2("Top-Right Controls")
bullet("Bell icon (🔔) —", "Opens the Live System Logs panel. A red dot means new system activity has been recorded.")
bullet("Theme toggle —", "Switch between light mode (white background) and dark mode (dark background). Choose whichever is easier on your eyes.")
bullet("Settings icon —", "Application settings (managed by your administrator).")

page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 3 — MAP VIEW
# ══════════════════════════════════════════════════════════════════════════════

heading1("3.  Map View — Your Main Workspace")

body(
    "The Map View is the heart of the platform. It shows the Athabasca watershed "
    "as an interactive map where you can see environmental hazard zones, click on "
    "monitoring sectors for details, and track live water quality readings."
)

heading2("What You See on the Map")
bullet("Base Map —", "A standard street and terrain map showing the Jasper Region and Athabasca River.")
bullet("Coloured Zones —", "Circles or highlighted areas representing environmental hazard zones. Each colour represents a different type of risk (see Layer Toggles below).")
bullet("Station Marker —", "A pin marking the central telemetry station (SEC-B4). This is a primary ground sensor location.")

heading2("Layer Toggles — Turning Map Layers On and Off")
body(
    "In the top-right corner of the map, you will find three switches. "
    "Each one controls a different layer of information on the map. "
    "Toggle them on or off depending on what you need to focus on."
)

layer_table = doc.add_table(rows=4, cols=3)
layer_table.style = "Table Grid"
for i, h in enumerate(["Layer", "Colour", "What it shows"]):
    cell = layer_table.cell(0, i)
    set_cell_bg(cell, "005F6E")
    p = cell.paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = WHITE
    r.font.size = Pt(11)

layers = [
    ("Erosion Outlines",   "Purple",    "Areas with slope erosion risk — where soil is unstable and may erode after rainfall."),
    ("Water Chemistry",    "Cyan/Blue", "Contaminant simulation zones showing how pollutants or sediments may spread through river flow."),
    ("Vegetation Index",   "Blue",      "Post-wildfire burn scar zones showing how much of the forest has been affected and the current recovery status."),
]
for i, (name, colour, desc) in enumerate(layers, start=1):
    layer_table.cell(i, 0).text = name
    layer_table.cell(i, 1).text = colour
    layer_table.cell(i, 2).text = desc
    layer_table.cell(i, 0).paragraphs[0].runs[0].bold = True

doc.add_paragraph()

tip_box(
    "If the map feels cluttered, turn off layers you are not currently using. "
    "For example, if you are assessing water contamination, turn off Erosion Outlines "
    "and Vegetation Index so only the Water Chemistry layer is visible."
)

heading2("Clicking on a Hazard Zone")
body(
    "You can click on any coloured zone on the map to open a details card. "
    "This card shows:"
)
bullet("The sector ID and name (e.g., SEC-E1 — Slope Sector).")
bullet("The current status: OPERATIONAL, WARNING, or CRITICAL.")
bullet("Key measurements for that zone, such as slope angle, water turbidity, or contamination level.")
bullet("The AI model's confidence score and risk label (Low / Medium / High).")

note_box(
    "The AI risk scores are produced by machine learning models trained on "
    "satellite imagery and sensor data. They are decision-support tools — "
    "always combine them with field observations and professional judgement."
)

heading2("Temporal Slider — Viewing Past and Present Conditions")
body(
    "At the bottom-left of the map is a time slider. Drag it left or right to "
    "step through different time periods — from the pre-fire baseline (June 2024) "
    "through to the current monitoring period. This lets you compare how conditions "
    "have changed over time."
)
bullet("Left side of slider —", "Pre-Fire conditions (baseline before the wildfire).")
bullet("Right side of slider —", "Most recent monitoring data.")

tip_box(
    "Use the temporal slider to show stakeholders or field teams the before-and-after "
    "difference in burn scar extent or water quality."
)

heading2("Sector Focus Panel")
body(
    "On the left sidebar, click 'Telemetry Focus' to see a list of all four monitoring "
    "sectors. Each one shows a colour-coded status badge:"
)

status_table = doc.add_table(rows=5, cols=3)
status_table.style = "Table Grid"
for i, h in enumerate(["Sector", "Status", "What it means"]):
    cell = status_table.cell(0, i)
    set_cell_bg(cell, "005F6E")
    p = cell.paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = WHITE
    r.font.size = Pt(11)

sectors = [
    ("SEC-B4 — Upper Stream",      "STABLE (Green)",    "All readings within normal range. No action required."),
    ("SEC-E1 — Slope Sector",      "HIGH (Red)",        "Elevated erosion risk detected. Review map zone and consider field inspection."),
    ("SEC-W2 — Athabasca Run",     "CRITICAL (Red)",    "Urgent conditions detected. Immediate review recommended."),
    ("SEC-V9 — Forest Segment",    "ELEVATED (Amber)",  "Conditions are above normal thresholds. Monitor closely."),
]
for i, (sector, status, meaning) in enumerate(sectors, start=1):
    status_table.cell(i, 0).text = sector
    status_table.cell(i, 1).text = status
    status_table.cell(i, 2).text = meaning

doc.add_paragraph()

body("Click 'Focus Map' next to any sector to automatically pan and zoom the map to that area.")

heading2("Live Telemetry Panel (Right Side of Map)")
body(
    "On the right side of the map screen, a panel shows live readings from the "
    "monitoring network. It updates automatically and contains four widgets:"
)
bullet("Water Quality —", "Current turbidity (water cloudiness) in NTU and pH level, with a small trend chart showing the last several readings.")
bullet("Pipeline Ingest —", "Shows how much satellite and IoT sensor data has been processed recently (shown as a percentage).")
bullet("Erosion Model —", "Displays the accuracy score of the erosion AI model and when it was last updated.")
bullet("Field Validation Photos —", "Placeholder tiles for field survey photos from ground teams.")

page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 4 — DASHBOARD
# ══════════════════════════════════════════════════════════════════════════════

heading1("4.  Dashboard — Water Quality Analytics")

body(
    "The Dashboard gives you a detailed view of water quality measurements across "
    "all monitoring stations. It is useful for trend analysis, preparing reports, "
    "or checking on specific stations after a weather event."
)

heading2("Selecting a Station")
body(
    "At the top-right of the Dashboard, you can filter by station using the buttons:"
)
bullet("All Stations —", "View a combined summary across the entire network.")
bullet("IoT Jasper-A1 —", "The primary IoT ground sensor near Jasper.")
bullet("Silt Monitor S-2 —", "A sediment and turbidity monitor on the river.")
bullet("Slope Sensor SL-4 —", "A stability sensor in a high-erosion area.")

heading2("The Four Metric Cards")
body(
    "Four summary cards display the most important water quality indicators. "
    "Click any card to see its detailed trend chart below."
)

metric_table = doc.add_table(rows=5, cols=3)
metric_table.style = "Table Grid"
for i, h in enumerate(["Metric", "What it measures", "Concern level"]):
    cell = metric_table.cell(0, i)
    set_cell_bg(cell, "005F6E")
    p = cell.paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = WHITE
    r.font.size = Pt(11)

metrics = [
    ("Turbidity",       "How cloudy or murky the water is. High turbidity can indicate sediment runoff or contamination.",        "Elevated = monitor; Severe = action required"),
    ("pH Level",        "The acidity or alkalinity of the water. Normal range is roughly 6.5–8.5.",                               "Outside 6.5–8.5 = investigate"),
    ("Ash Concentration", "Level of ash or particulate matter in the water — a direct indicator of wildfire impact.",             "High Risk = urgent review"),
    ("Soil Stability",  "A percentage estimate of how stable the surrounding soil is. Lower values mean higher erosion risk.",    "Below 80% = elevated concern"),
]
for i, (m, desc, concern) in enumerate(metrics, start=1):
    metric_table.cell(i, 0).text = m
    metric_table.cell(i, 1).text = desc
    metric_table.cell(i, 2).text = concern
    metric_table.cell(i, 0).paragraphs[0].runs[0].bold = True

doc.add_paragraph()

heading2("The Trend Chart")
body(
    "The large chart on the left shows hourly readings for the selected metric "
    "over the past 12 hours. A dashed orange line marks the 'Danger Baseline' — "
    "the threshold above (or below) which conditions become concerning. If the blue "
    "line crosses the orange line, pay attention to the corresponding alert in the "
    "panel on the right."
)

heading2("Watershed Status Alerts (Right Panel)")
body("The right side of the Dashboard shows three persistent alert items:")
bullet("Active Runoff Risk Warning (Red) —", "Soil hydration levels are high and erosion potential is elevated in burn scar areas.")
bullet("Turbidity Surge Detected (Amber) —", "Water cloudiness is rising, possibly due to rainfall or upstream sediment movement.")
bullet("Sentinel-2 Ingest Stable (Green) —", "Satellite data is being processed successfully. No action needed.")

tip_box(
    "A green alert is good news — it means the data pipeline is healthy and the "
    "readings you are seeing are up to date."
)

page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 5 — REPORTS
# ══════════════════════════════════════════════════════════════════════════════

heading1("5.  Reports — Download & Generate Reports")

body(
    "The Reports section is your library of environmental diagnostic documents. "
    "You can download completed reports as PDF files, or trigger new model runs "
    "to generate updated assessments."
)

heading2("Browsing Reports")
body(
    "Reports are listed as cards. Each card shows the report title, the author or "
    "organization, the publication date, file size, and its current status."
)

heading2("Report Statuses")
status_rep_table = doc.add_table(rows=4, cols=2)
status_rep_table.style = "Table Grid"
for i, h in enumerate(["Status", "What it means"]):
    cell = status_rep_table.cell(0, i)
    set_cell_bg(cell, "005F6E")
    p = cell.paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = WHITE
    r.font.size = Pt(11)

rep_statuses = [
    ("Ready (Green ✓)",        "The report is complete and ready to download. Click 'Download PDF Report' to save it to your computer."),
    ("Draft Model (Clock)",    "The model has been set up but not yet run. Click 'Run Stability Model Ingest' to start generating the report."),
    ("Running... (Clock)",     "The model is currently processing. This may take a few minutes. Refresh the page to check if it has completed."),
]
for i, (s, m) in enumerate(rep_statuses, start=1):
    status_rep_table.cell(i, 0).text = s
    status_rep_table.cell(i, 1).text = m
    status_rep_table.cell(i, 0).paragraphs[0].runs[0].bold = True

doc.add_paragraph()

heading2("Filtering Reports by Category")
body(
    "Use the category filter buttons at the top-right of the Reports page to narrow "
    "down the list:"
)
bullet("All —", "Show every report.")
bullet("Hydro-geology —", "Reports about water flow, soil, and geological conditions.")
bullet("Limnology —", "Reports about lake and river water quality.")
bullet("Remote Sensing —", "Reports based on satellite imagery analysis.")
bullet("Model Output —", "Reports generated by the AI simulation models.")

heading2("Available Reports")
body("The following reports are currently in the system:")
bullet("Jasper Post-Fire Runoff & Erosion Risk Model —", "Prepared by Dr. Eleanor Vance. Covers erosion risk across burn scar areas.")
bullet("Athabasca Watershed Basin Water Quality Diagnostic —", "Prepared by the Jasper GIS Alpha Team. Full water quality assessment.")
bullet("Sentinel-2 Multi-Spectral Biomass Recovery Analysis —", "Prepared by the Canadian Forestry Service. Vegetation recovery from satellite data.")
bullet("Soil Stability Assessment & Contour Plume Prediction —", "Generated by Erosion Predictor v2.4. AI model output for soil stability.")

note_box(
    "Downloaded reports are saved to your computer's default Downloads folder "
    "unless your browser is configured differently."
)

page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 6 — ARCHIVES
# ══════════════════════════════════════════════════════════════════════════════

heading1("6.  Archives — Historical Survey Data")

body(
    "The Archives section stores historical environmental surveys and satellite "
    "passes. Use it to look back at how conditions were before and after the "
    "wildfire, or to compare current readings against a historical baseline."
)

heading2("Searching the Archives")
body(
    "Use the search bar in the top-right of the Archives page to find a specific "
    "record by name or archive ID (e.g., type 'Pre-Fire' or 'ARC-014')."
)

heading2("Archive Types")
arch_table = doc.add_table(rows=4, cols=2)
arch_table.style = "Table Grid"
for i, h in enumerate(["Type", "What it contains"]):
    cell = arch_table.cell(0, i)
    set_cell_bg(cell, "005F6E")
    p = cell.paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = WHITE
    r.font.size = Pt(11)

arch_types = [
    ("Full Survey (Blue)",        "A complete assessment of the entire watershed, including water quality, soil, and vegetation data."),
    ("Incremental (Purple)",      "A partial update covering only the changes since the last full survey."),
    ("Satellite Pass (Cyan)",     "Raw and processed data from a Sentinel-2 satellite fly-over of the region."),
]
for i, (t, d) in enumerate(arch_types, start=1):
    arch_table.cell(i, 0).text = t
    arch_table.cell(i, 1).text = d
    arch_table.cell(i, 0).paragraphs[0].runs[0].bold = True

doc.add_paragraph()

heading2("Viewing an Archive Record")
body(
    "Click the 'View Snapshot' button on any archive record to expand it and see "
    "a summary of its contents, the coverage area, and any associated notes."
)

heading2("Available Archives")
body("The following historical records are currently in the system:")

archive_list = [
    ("ARC-014", "Pre-Fire Baseline Survey",         "May 2023",  "212 MB", "Full Survey"),
    ("ARC-015", "Post-Fire Initial Assessment",     "Aug 2023",  "248 MB", "Full Survey"),
    ("ARC-016", "Q4 2023 Recovery Snapshot",        "Dec 2023",  "64 MB",  "Incremental"),
    ("ARC-017", "Sentinel-2 Winter Pass",           "Feb 2024",  "97 MB",  "Satellite Pass"),
    ("ARC-018", "Q2 2024 Recovery Snapshot",        "May 2024",  "71 MB",  "Incremental"),
    ("ARC-019", "Mid-Summer Watershed Survey",      "Jul 2024",  "256 MB", "Full Survey"),
]
arc_table = doc.add_table(rows=len(archive_list)+1, cols=4)
arc_table.style = "Table Grid"
for i, h in enumerate(["ID", "Name", "Date", "Type"]):
    cell = arc_table.cell(0, i)
    set_cell_bg(cell, "005F6E")
    p = cell.paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = WHITE
    r.font.size = Pt(11)

for i, (aid, name, date, size, atype) in enumerate(archive_list, start=1):
    arc_table.cell(i, 0).text = aid
    arc_table.cell(i, 1).text = name
    arc_table.cell(i, 2).text = date
    arc_table.cell(i, 3).text = atype

doc.add_paragraph()

page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 7 — ALERTS
# ══════════════════════════════════════════════════════════════════════════════

heading1("7.  Alerts & What They Mean")

body(
    "The platform uses a consistent colour-coded alert system across the Map View, "
    "Dashboard, and sector panels. Here is what each level means and what action "
    "to consider."
)

alert_table = doc.add_table(rows=5, cols=3)
alert_table.style = "Table Grid"
for i, h in enumerate(["Colour / Label", "What it means", "Suggested action"]):
    cell = alert_table.cell(0, i)
    set_cell_bg(cell, "005F6E")
    p = cell.paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = WHITE
    r.font.size = Pt(11)

alerts = [
    ("🟢 Green — STABLE / OPERATIONAL / Optimal",
     "All readings are within normal ranges. The system and sensors are working correctly.",
     "No action required. Continue routine monitoring."),
    ("🟡 Amber — ELEVATED / WARNING",
     "One or more readings are above normal thresholds but not yet at a critical level.",
     "Monitor more frequently. Consider field inspection if the condition persists."),
    ("🔴 Red — HIGH / CRITICAL / High Risk",
     "Readings are significantly outside safe ranges. Conditions may pose an environmental or safety risk.",
     "Review immediately. Notify relevant team members. Consider field response."),
    ("⚪ Grey — Neutral / Nominal / INFO",
     "Normal system activity or informational status. No concern indicated.",
     "No action required."),
]
for i, (label, meaning, action) in enumerate(alerts, start=1):
    alert_table.cell(i, 0).text = label
    alert_table.cell(i, 1).text = meaning
    alert_table.cell(i, 2).text = action

doc.add_paragraph()

note_box(
    "Alert levels are set based on thresholds defined by your environmental team. "
    "If you believe a threshold needs adjusting, contact your system administrator."
)

page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 8 — LIVE SYSTEM LOGS
# ══════════════════════════════════════════════════════════════════════════════

heading1("8.  Live System Logs")

body(
    "The Live GIS Logs panel is a real-time feed of system activity. It shows you "
    "what the platform is doing behind the scenes — ingesting satellite data, "
    "running models, polling sensors, and flagging any errors."
)
body(
    "To open it, click the bell icon (🔔) in the top-right corner of any screen, "
    "or click 'Diagnostic Logs' at the bottom of the left sidebar."
)

heading2("Reading the Logs")
body("Each log entry has three parts:")
bullet("Timestamp —", "The time the event occurred (24-hour format, e.g., 14:32:05).")
bullet("Level badge —", "Colour-coded category (see table below).")
bullet("Message —", "A plain-English description of what happened.")

log_table = doc.add_table(rows=5, cols=2)
log_table.style = "Table Grid"
for i, h in enumerate(["Badge", "What it means"]):
    cell = log_table.cell(0, i)
    set_cell_bg(cell, "005F6E")
    p = cell.paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = WHITE
    r.font.size = Pt(11)

log_levels = [
    ("INFO (Grey)",    "Routine system activity. The platform is working normally."),
    ("SUCCESS (Green)","An operation completed successfully (e.g., satellite data synced)."),
    ("WARN (Amber)",   "Something unexpected happened but the system recovered. Worth noting."),
    ("ERROR (Red)",    "Something failed. The system will usually retry automatically, but persistent errors should be reported."),
]
for i, (badge, meaning) in enumerate(log_levels, start=1):
    log_table.cell(i, 0).text = badge
    log_table.cell(i, 1).text = meaning
    log_table.cell(i, 0).paragraphs[0].runs[0].bold = True

doc.add_paragraph()

heading2("Log Panel Controls")
bullet("Pause / Resume Ingest —", "Temporarily stops new log entries from appearing so you can read the current ones.")
bullet("Clear —", "Removes all visible log entries from the panel (does not delete any data).")
bullet("Export Logs —", "Downloads the visible log entries as a file for record-keeping or sharing with the support team.")

tip_box(
    "If you see repeated ERROR entries for the same sensor or station, take note of "
    "the station name and time, and report it using the Support Request link in the sidebar."
)

page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 9 — GLOSSARY
# ══════════════════════════════════════════════════════════════════════════════

heading1("9.  Glossary of Terms")

body(
    "Below is a plain-English explanation of the technical terms you may encounter "
    "while using the platform."
)

terms = [
    ("Ash Concentration",   "The amount of ash or burned particulate matter detected in the water, measured in parts per million (ppm). A sign of direct wildfire impact on the water supply."),
    ("Burn Scar",           "An area of land where vegetation was destroyed by the wildfire. Burn scars are mapped from satellite imagery."),
    ("Contaminant Plume",   "A simulated representation of how a pollutant or sediment might spread through water, based on flow direction and speed."),
    ("Confidence Score",    "A percentage or 0–1 value indicating how certain the AI model is about its prediction. Higher is more reliable."),
    ("GeoJSON",             "A file format for geographic data. Used by GIS software like ArcGIS or QGIS."),
    ("IoT Station",         "Internet of Things station — a ground-based sensor that automatically measures and transmits environmental data (water quality, slope stability, etc.)."),
    ("NTU",                 "Nephelometric Turbidity Unit — the standard measurement for water cloudiness. Pure water is near 0 NTU; highly turbid water can exceed 100 NTU."),
    ("pH",                  "A scale from 0–14 measuring how acidic or alkaline water is. 7 is neutral; below 7 is acidic; above 7 is alkaline. Healthy river water is typically 6.5–8.5."),
    ("PostGIS",             "The database technology used to store and query the geographic (spatial) data in this platform."),
    ("ppm",                 "Parts per million — a unit of concentration used for measuring trace substances like ash or chemicals in water."),
    ("Risk Label",          "The AI model's classification of hazard severity: Low, Medium, or High."),
    ("RUSLE",               "Revised Universal Soil Loss Equation — a scientific formula used by the erosion model to estimate how much soil may erode based on slope, rainfall, and vegetation cover."),
    ("Sector",              "A defined geographic sub-area of the watershed used to organize monitoring data. Each sector has its own sensors and status."),
    ("Sentinel-2",          "A European Space Agency satellite that provides free, regular satellite imagery of Earth's surface. Project Jasper uses it for vegetation and land cover analysis."),
    ("Telemetry",           "Automated data collected and transmitted from remote sensors without manual intervention."),
    ("Turbidity",           "A measure of water cloudiness caused by suspended particles such as sediment, ash, or algae. High turbidity can be harmful to aquatic life and indicates water quality issues."),
]

gloss_table = doc.add_table(rows=len(terms)+1, cols=2)
gloss_table.style = "Table Grid"
for i, h in enumerate(["Term", "Definition"]):
    cell = gloss_table.cell(0, i)
    set_cell_bg(cell, "005F6E")
    p = cell.paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = WHITE
    r.font.size = Pt(11)

for i, (term, defn) in enumerate(terms, start=1):
    gloss_table.cell(i, 0).text = term
    gloss_table.cell(i, 1).text = defn
    gloss_table.cell(i, 0).paragraphs[0].runs[0].bold = True

doc.add_paragraph()
page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 10 — FAQ
# ══════════════════════════════════════════════════════════════════════════════

heading1("10.  Frequently Asked Questions")

faqs = [
    (
        "The map is not loading or appears blank. What should I do?",
        "First, refresh your browser (F5 or Ctrl+R). If the map still does not appear, "
        "check your internet connection. The map tiles are loaded from an external "
        "service — a slow connection can delay their appearance. If the issue persists "
        "for more than a few minutes, contact support."
    ),
    (
        "The data on the Dashboard looks old. Is it up to date?",
        "The dashboard data refreshes automatically. If you suspect it is stale, "
        "check the Live System Logs panel for any WARN or ERROR entries about "
        "sensor connectivity. A green 'Sentinel-2 Ingest Stable' alert on the "
        "Dashboard confirms that satellite data is flowing normally."
    ),
    (
        "A sector is showing CRITICAL status. What do I do?",
        "Click on the sector in the Telemetry Focus panel or on the map to see the "
        "detailed readings. Review which specific measurements are out of range. "
        "Follow your organization's standard operating procedure for environmental "
        "alerts — this may involve notifying field teams or a senior scientist."
    ),
    (
        "How do I download a report?",
        "Go to the Reports tab. Find the report you need and click the blue "
        "'Download PDF Report' button. The file will be saved to your browser's "
        "default download location."
    ),
    (
        "What does 'Run Stability Model Ingest' do?",
        "It starts a new AI model run to generate an updated version of that report. "
        "The status will change to 'Running...' while the model processes. Once "
        "complete, the status will show 'Ready' and you can download the updated PDF."
    ),
    (
        "Can I export the map data to use in ArcGIS or another GIS tool?",
        "Yes. Click 'Export GeoJSON' at the bottom of the left sidebar. This downloads "
        "the current map layer data in GeoJSON format, which is compatible with most "
        "professional GIS software."
    ),
    (
        "I see a lot of ERROR messages in the logs. Is this serious?",
        "Occasional errors are normal — sensors can briefly lose connectivity and "
        "recover automatically. If you see the same ERROR repeating every few minutes "
        "for the same station, that station may need attention. Use 'Export Logs' to "
        "save the log entries and contact support with the details."
    ),
    (
        "Can I switch between light and dark mode?",
        "Yes — click the theme toggle icon in the top-right corner of any screen "
        "to switch between light and dark mode. Your preference is not saved between "
        "sessions, so you may need to switch each time you open the application."
    ),
]

for question, answer in faqs:
    p_q = doc.add_paragraph()
    p_q.paragraph_format.space_before = Pt(10)
    p_q.paragraph_format.space_after  = Pt(2)
    rq = p_q.add_run("Q:  " + question)
    rq.bold = True
    rq.font.size = Pt(11)
    rq.font.color.rgb = TEAL

    p_a = doc.add_paragraph()
    p_a.paragraph_format.space_after = Pt(8)
    ra = p_a.add_run("A:  " + answer)
    ra.font.size = Pt(11)
    ra.font.color.rgb = DARK

page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 11 — CONTACT & SUPPORT
# ══════════════════════════════════════════════════════════════════════════════

heading1("11.  Contact & Support")

body(
    "If you encounter an issue that is not covered in this guide, or if you need "
    "to request a new feature or report a data problem, please use one of the "
    "following channels."
)

heading2("In-App Support")
body(
    "Click 'Support Request' at the bottom of the left sidebar. This will open "
    "a contact form pre-filled with your session information to help the technical "
    "team diagnose the issue quickly."
)

heading2("Providing Useful Information When Reporting an Issue")
body("When contacting support, please include:")
bullet("The page or feature where the issue occurred (e.g., Map View, Dashboard, Reports).")
bullet("What you expected to happen.")
bullet("What actually happened.")
bullet("Any error messages you saw (use 'Export Logs' in the Live System Logs panel to capture them).")
bullet("The date and approximate time the issue occurred.")

heading2("Project Team")
team_table = doc.add_table(rows=6, cols=3)
team_table.style = "Table Grid"
for i, h in enumerate(["Name", "Role", "Area"]):
    cell = team_table.cell(0, i)
    set_cell_bg(cell, "005F6E")
    p = cell.paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = WHITE
    r.font.size = Pt(11)

team = [
    ("Edwin Olaez",    "Project Manager & QA/Security Lead", "CI/CD, testing, documentation"),
    ("Feven Ytbarek",  "Data Pipeline & API",                "Backend, ingest pipeline"),
    ("Richard Li",     "AI/ML & Simulation",                 "ML models, simulation endpoints"),
    ("Reyta",          "Frontend & GIS",                     "Map UI, dashboard, Next.js"),
    ("Rahil Khan",     "Database & Analytics",               "Supabase, PostGIS, Convex"),
]
for i, (name, role, area) in enumerate(team, start=1):
    team_table.cell(i, 0).text = name
    team_table.cell(i, 1).text = role
    team_table.cell(i, 2).text = area

doc.add_paragraph()

divider()

footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = footer.add_run(
    f"Project Jasper User Guide  ·  Version 1.0  ·  {datetime.date.today().strftime('%B %Y')}\n"
    "Prepared by Team BluePulse AI — SAIT Capstone 2026\n"
    "For CERCUTS / ARIS — Athabasca Watershed Environmental Monitoring"
)
rf.font.size = Pt(9)
rf.font.color.rgb = GREY
rf.italic = True


# ── Save ──────────────────────────────────────────────────────────────────────
output_path = r"C:\Users\Edwin Olaez\Documents\BluePulseAi-Capstone\docs\Project-Jasper-User-Guide.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
